import docx


def create_invitations(guests_file, output_file):
    # Create a new Word document
    doc = docx.Document('guests.txt')

    # Add custom styles for the invitation
    styles = doc.styles
    invitation_style = styles.add_style('Invitation', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    invitation_style.font.name = 'Arial'
    invitation_style.font.size = docx.shared.Pt(12)
    invitation_style.font.bold = True
    invitation_style.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

    # Read guest names from the file and create invitations
    with open(guests_file, 'r') as file:
        guest_names = file.read().splitlines()

    for guest_name in guest_names:
        # Add invitation text with custom style
        doc.add_paragraph(f"Dear {guest_name},", style='Invitation')
        doc.add_paragraph("We cordially invite you to our event.", style='Invitation')
        doc.add_paragraph("Date: [Event Date]", style='Invitation')
        doc.add_paragraph("Location: [Event Location]", style='Invitation')
        doc.add_paragraph("Please RSVP by [RSVP Date].", style='Invitation')

        # Add a page break after each invitation
        doc.add_page_break()

    # Save the Word document
    doc.save(output_file)


if __name__ == "__main__":
    guests_file = "guests.txt"  # Replace with your actual file path
    output_file = "invitations.docx"  # Replace with your desired output file path

    create_invitations(guests_file, output_file)
